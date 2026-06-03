"""Exportador de FDS a PDF (WeasyPrint), JSON, CSV y DOCX."""

import csv
import io
import json
import os
from datetime import datetime, timezone

from config import Config
from traductor_fds import traducir_fds


def exportar_json(fds_secciones, clasificacion, mezcla_data):
    """Exporta la FDS completa a formato JSON."""
    return json.dumps({
        "mezcla": mezcla_data,
        "clasificacion": clasificacion,
        "secciones": fds_secciones,
        "fecha_exportacion": datetime.now(timezone.utc).isoformat(),
    }, ensure_ascii=False, indent=2)


def exportar_csv(fds_secciones, clasificacion, mezcla_data):
    """Exporta la FDS a formato CSV (tabla plana para auditorias)."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Seccion", "Titulo", "Contenido"])

    for num in sorted(fds_secciones.keys(), key=lambda x: int(x)):
        sec = fds_secciones[num]
        titulo = sec.get("titulo", f"Seccion {num}")
        contenido = sec.get("contenido", "")
        # Limpiar saltos de linea para CSV
        contenido_limpio = contenido.replace("\n", " | ")
        writer.writerow([num, titulo, contenido_limpio])

    # Agregar fila de clasificacion
    writer.writerow([])
    writer.writerow(["Clasificacion GHS"])
    writer.writerow(["Codigos H", ", ".join(clasificacion.get("codigos_h", []))])
    writer.writerow(["Pictogramas", ", ".join(clasificacion.get("pictogramas", []))])
    writer.writerow(["Palabra de advertencia", clasificacion.get("palabra_advertencia", "")])

    return output.getvalue()


def _generar_html_fds(fds_secciones, clasificacion, mezcla_data):
    """Genera HTML completo de la FDS para renderizar a PDF."""
    nombre = mezcla_data.get("nombre_producto", "Sin nombre")
    lang = mezcla_data.get("_lang", "es")
    is_en = lang == "en"
    header_title = "SAFETY DATA SHEET" if is_en else "FICHA DE DATOS DE SEGURIDAD"
    header_label = "SAFETY DATA SHEET" if is_en else "FICHA DE DATOS DE SEGURIDAD"
    page_label = "Page" if is_en else "Pagina"
    page_of = "of" if is_en else "de"
    pictogramas = clasificacion.get("pictogramas", [])
    pictos_dir = os.path.join(Config.BASE_DIR, "static", "pictogramas")

    pictos_html = ""
    for p in pictogramas:
        svg_path = os.path.join(pictos_dir, f"{p}.svg")
        if os.path.exists(svg_path):
            with open(svg_path, "r") as f:
                pictos_html += f'<div class="pictograma">{f.read()}</div>'

    secciones_html = ""
    for num in sorted(fds_secciones.keys(), key=lambda x: int(x)):
        sec = fds_secciones[num]
        titulo = sec.get("titulo", f"Seccion {num}")
        contenido = sec.get("contenido", "").replace("\n", "<br>")
        secciones_html += f"""
        <div class="seccion">
            <h2>{titulo}</h2>
            <div class="contenido">{contenido}</div>
        </div>
        """

    html_lang = "en" if is_en else "es"
    html = f"""<!DOCTYPE html>
<html lang="{html_lang}">
<head>
    <meta charset="UTF-8">
    <title>{"SDS" if is_en else "FDS"} - {nombre}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
            @top-center {{
                content: "{header_label} - {nombre}";
                font-size: 8pt;
                color: #666;
            }}
            @bottom-center {{
                content: "{page_label} " counter(page) " {page_of} " counter(pages);
                font-size: 8pt;
                color: #666;
            }}
        }}
        body {{
            font-family: Arial, Helvetica, sans-serif;
            font-size: 10pt;
            line-height: 1.4;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 2px solid #c00;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .header h1 {{
            color: #c00;
            font-size: 16pt;
            margin: 0;
        }}
        .pictogramas {{
            display: flex;
            justify-content: center;
            gap: 10px;
            margin: 10px 0;
        }}
        .pictograma {{
            width: 60px;
            height: 60px;
        }}
        .pictograma svg {{
            width: 100%;
            height: 100%;
        }}
        .seccion {{
            margin-bottom: 15px;
            page-break-inside: avoid;
        }}
        .seccion h2 {{
            color: #c00;
            font-size: 11pt;
            border-bottom: 1px solid #ddd;
            padding-bottom: 3px;
            margin-top: 15px;
            margin-bottom: 5px;
        }}
        .contenido {{
            margin-left: 10px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{header_title}</h1>
        <h2>{nombre}</h2>
        <div class="pictogramas">{pictos_html}</div>
    </div>
    {secciones_html}
</body>
</html>"""
    return html


def exportar_pdf(fds_secciones, clasificacion, mezcla_data, ruta_salida=None):
    """Exporta la FDS a PDF usando WeasyPrint.

    Args:
        fds_secciones: dict con secciones
        clasificacion: dict con clasificacion GHS
        mezcla_data: dict con datos de la mezcla
        ruta_salida: ruta del archivo PDF de salida (opcional)

    Returns:
        bytes del PDF si no se especifica ruta, o la ruta del archivo creado
    """
    try:
        from weasyprint import HTML
    except ImportError:
        raise RuntimeError(
            "WeasyPrint no esta instalado. Ejecute: "
            "sudo apt install -y libpango-1.0-0 libcairo2 libgdk-pixbuf2.0-0 libffi-dev shared-mime-info && "
            "pip install weasyprint"
        )

    html = _generar_html_fds(fds_secciones, clasificacion, mezcla_data)

    if ruta_salida:
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        HTML(string=html).write_pdf(ruta_salida)
        return ruta_salida
    else:
        return HTML(string=html).write_pdf()


def exportar_pdf_en(fds_secciones, clasificacion, mezcla_data, ruta_salida=None):
    """Exporta la FDS en inglés (SDS) a PDF."""
    secciones_en = traducir_fds(fds_secciones)
    mezcla_en = dict(mezcla_data)
    mezcla_en["_lang"] = "en"
    return exportar_pdf(secciones_en, clasificacion, mezcla_en, ruta_salida)


def exportar_docx(fds_secciones, clasificacion, mezcla_data):
    """Exporta la FDS a formato DOCX (Word).

    Returns:
        bytes del archivo DOCX
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx no esta instalado. Ejecute: pip install python-docx"
        )

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    nombre = mezcla_data.get("nombre_producto", "Sin nombre")
    lang = mezcla_data.get("_lang", "es")
    is_en = lang == "en"

    # Header
    header_title = "SAFETY DATA SHEET" if is_en else "FICHA DE DATOS DE SEGURIDAD"
    heading = doc.add_heading(header_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    subtitle = doc.add_paragraph(nombre)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(14)

    # Pictograms line
    pictogramas = clasificacion.get("pictogramas", [])
    if pictogramas:
        pictos_label = "Pictograms" if is_en else "Pictogramas"
        p = doc.add_paragraph(f"{pictos_label}: {', '.join(pictogramas)}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    palabra_adv = clasificacion.get("palabra_advertencia", "")
    if palabra_adv:
        p = doc.add_paragraph(palabra_adv)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph("")  # spacer

    # Sections
    for num in sorted(fds_secciones.keys(), key=lambda x: int(x)):
        sec = fds_secciones[num]
        titulo = sec.get("titulo", f"Section {num}" if is_en else f"Seccion {num}")
        contenido = sec.get("contenido", "")

        heading = doc.add_heading(titulo, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.size = Pt(11)

        for line in contenido.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    # Footer
    doc.add_paragraph("")
    footer_text = (
        "Generated by Sistema FSD 2026"
        if is_en
        else "Generado por Sistema FSD 2026"
    )
    p = doc.add_paragraph(footer_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def exportar_docx_en(fds_secciones, clasificacion, mezcla_data):
    """Exporta la FDS en inglés (SDS) a DOCX."""
    secciones_en = traducir_fds(fds_secciones)
    mezcla_en = dict(mezcla_data)
    mezcla_en["_lang"] = "en"
    return exportar_docx(secciones_en, clasificacion, mezcla_en)
